from __future__ import annotations

from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path("dp_experiments/outputs/complete_original_models")
REPORT_DIR = ROOT / "report"
FIG_DIR = REPORT_DIR / "figures"
DOCX_PATH = REPORT_DIR / "DP_synthetic_training_top_journal_report.docx"


METHOD_LABELS = {
    "dp": "DP",
    "gan_dp": "GAN+DP",
    "distill_dp": "Dataset Distillation+DP",
    "diffusion_dp": "Diffusion+DP",
}


def load_data() -> pd.DataFrame:
    frames = []
    for dataset in ["mnist", "cifar10"]:
        df = pd.read_csv(ROOT / dataset / "metrics.csv")
        df.insert(0, "dataset", dataset.upper() if dataset == "mnist" else "CIFAR10")
        frames.append(df)
    data = pd.concat(frames, ignore_index=True)
    data["method_label"] = data["method"].map(METHOD_LABELS)
    return data


def ensure_dirs() -> None:
    FIG_DIR.mkdir(parents=True, exist_ok=True)


def setup_plot_style() -> None:
    plt.rcParams["figure.dpi"] = 220
    plt.rcParams["savefig.dpi"] = 300
    plt.rcParams["font.family"] = "DejaVu Sans"
    plt.rcParams["axes.grid"] = True
    plt.rcParams["grid.alpha"] = 0.25
    plt.rcParams["axes.spines.top"] = False
    plt.rcParams["axes.spines.right"] = False


def save_fig(fig, name: str) -> Path:
    path = FIG_DIR / f"{name}.png"
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


def make_figures(data: pd.DataFrame) -> dict[str, Path]:
    figs: dict[str, Path] = {}

    best = (
        data.sort_values("original_acc", ascending=False)
        .groupby("dataset", as_index=False)
        .first()
    )
    fig, ax = plt.subplots(figsize=(6.6, 4.0))
    labels = [
        f"{r.dataset}\n{METHOD_LABELS[r.method]}\neps={r.epsilon}, k={int(r.per_class)}, {r.classifier}"
        for r in best.itertuples()
    ]
    ax.bar(labels, best["original_acc"] * 100, color=["#2563EB", "#10B981"])
    ax.set_ylabel("Original validation accuracy (%)")
    ax.set_title("Best Configuration on Each Dataset")
    ax.set_ylim(0, 105)
    for i, v in enumerate(best["original_acc"] * 100):
        ax.text(i, v + 1, f"{v:.2f}%", ha="center", fontsize=10)
    figs["best"] = save_fig(fig, "fig1_best_config")

    method_mean = (
        data.groupby(["dataset", "method_label"], as_index=False)["original_acc"]
        .mean()
        .sort_values(["dataset", "original_acc"], ascending=[True, False])
    )
    fig, axes = plt.subplots(1, 2, figsize=(9.8, 4.0), sharey=True)
    for ax, dataset in zip(axes, ["MNIST", "CIFAR10"]):
        sub = method_mean[method_mean["dataset"] == dataset]
        ax.bar(sub["method_label"], sub["original_acc"] * 100, color="#4F46E5")
        ax.set_title(dataset)
        ax.tick_params(axis="x", rotation=25)
        ax.set_ylabel("Mean original accuracy (%)")
        for i, v in enumerate(sub["original_acc"] * 100):
            ax.text(i, v + 1, f"{v:.1f}", ha="center", fontsize=8)
    figs["method"] = save_fig(fig, "fig2_method_mean")

    eps_mean = data.groupby(["dataset", "epsilon", "method_label"], as_index=False)[
        "original_acc"
    ].mean()
    fig, axes = plt.subplots(1, 2, figsize=(10.2, 4.0), sharey=False)
    for ax, dataset in zip(axes, ["MNIST", "CIFAR10"]):
        sub = eps_mean[eps_mean["dataset"] == dataset]
        for method, msub in sub.groupby("method_label"):
            ax.plot(msub["epsilon"], msub["original_acc"] * 100, marker="o", label=method)
        ax.set_title(f"{dataset}: epsilon vs utility")
        ax.set_xlabel("epsilon")
        ax.set_ylabel("Mean original accuracy (%)")
        ax.legend(fontsize=7)
    figs["epsilon"] = save_fig(fig, "fig3_epsilon_trend")

    k_mean = data.groupby(["dataset", "per_class", "method_label"], as_index=False)[
        "original_acc"
    ].mean()
    fig, axes = plt.subplots(1, 2, figsize=(10.2, 4.0), sharey=False)
    for ax, dataset in zip(axes, ["MNIST", "CIFAR10"]):
        sub = k_mean[k_mean["dataset"] == dataset]
        for method, msub in sub.groupby("method_label"):
            ax.plot(msub["per_class"], msub["original_acc"] * 100, marker="s", label=method)
        ax.set_title(f"{dataset}: synthetic images per class")
        ax.set_xlabel("k")
        ax.set_ylabel("Mean original accuracy (%)")
        ax.legend(fontsize=7)
    figs["k"] = save_fig(fig, "fig4_k_trend")

    time_mean = (
        data.groupby(["dataset", "method_label"], as_index=False)[
            ["synthesis_seconds", "train_seconds"]
        ]
        .mean()
        .melt(
            id_vars=["dataset", "method_label"],
            value_vars=["synthesis_seconds", "train_seconds"],
            var_name="time_type",
            value_name="seconds",
        )
    )
    fig, axes = plt.subplots(1, 2, figsize=(10.4, 4.0), sharey=False)
    for ax, dataset in zip(axes, ["MNIST", "CIFAR10"]):
        sub = time_mean[time_mean["dataset"] == dataset]
        xlabels = list(sub["method_label"].unique())
        x = range(len(xlabels))
        synth = [
            sub[(sub["method_label"] == m) & (sub["time_type"] == "synthesis_seconds")][
                "seconds"
            ].iloc[0]
            for m in xlabels
        ]
        train = [
            sub[(sub["method_label"] == m) & (sub["time_type"] == "train_seconds")][
                "seconds"
            ].iloc[0]
            for m in xlabels
        ]
        ax.bar([i - 0.18 for i in x], synth, width=0.36, label="synthesis")
        ax.bar([i + 0.18 for i in x], train, width=0.36, label="training")
        ax.set_xticks(list(x))
        ax.set_xticklabels(xlabels, rotation=25)
        ax.set_title(f"{dataset}: mean runtime")
        ax.set_ylabel("seconds")
        ax.legend(fontsize=8)
    figs["time"] = save_fig(fig, "fig5_time_comparison")

    return figs


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(9)


def add_table(doc: Document, rows, headers) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val)
    doc.add_paragraph()


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")


def set_doc_style(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.3)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    styles["Normal"].font.size = Pt(10.5)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.name = "Times New Roman"
        styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")


def add_page_number(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def p(doc: Document, text: str, bold: bool = False) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(10.5)


def pct(x) -> str:
    return f"{float(x) * 100:.2f}%"


def sec(x) -> str:
    return f"{float(x):.2f}"


def build_doc(data: pd.DataFrame, figs: dict[str, Path]) -> None:
    doc = Document()
    set_doc_style(doc)
    add_page_number(doc.sections[0])

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("本地差分隐私合成数据外包训练实验：MNIST 与 CIFAR10 的隐私-效用权衡评估")
    r.bold = True
    r.font.size = Pt(16)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Top-journal style experimental report")
    r.italic = True
    r.font.size = Pt(10)

    doc.add_paragraph()
    doc.add_heading("摘要", level=1)
    p(
        doc,
        "本文围绕企业本地数据不外泄、利用第三方算力训练模型的应用场景，"
        "系统评估了四类本地脱敏/合成方法（DP、GAN+DP、Dataset Distillation+DP、Diffusion+DP）"
        "在 MNIST 与 CIFAR10 上的隐私预算、合成样本规模、训练耗时和下游模型精度之间的权衡。"
        "实验保持下游模型结构不变，使用 CNN 与 ResNet20 作为第三方训练模型，"
        "以原始验证集准确率作为主要效用指标，以合成验证集准确率、生成时间、训练时间和视觉近邻对比作为辅助指标。"
        "最新完整实验显示，MNIST 上 Diffusion+DP 在 epsilon=0.5、k=320、ResNet20 下取得 96.91% 的最高原始验证精度；"
        "CIFAR10 上 Diffusion+DP 在 epsilon=1.0、k=320、CNN 下取得 47.66% 的最高原始验证精度。"
        "整体上，Diffusion+DP 在两类数据集上均表现出最佳平均效用，扩大每类合成样本数 k 是当前最稳定的提精度方向。",
    )
    p(doc, "关键词：差分隐私；合成数据；外包训练；MNIST；CIFAR10；隐私-效用权衡")

    doc.add_heading("1. 任务图逐项回应", level=1)
    rows = [
        [
            "企业本地数据不泄露给远程服务器",
            "本实验只将本地脱敏/合成后的数据用于第三方训练流程，原始数据仅在本地用于生成合成数据与原始验证评估。",
        ],
        [
            "对 MNIST、CIFAR10 进行脱敏处理",
            "已完成。两个数据集均运行完整矩阵，结果分别保存在 mnist/metrics.csv 与 cifar10/metrics.csv。",
        ],
        [
            "生成 4 类脱敏/合成数据集",
            "已完成 DP、GAN+DP、Dataset Distillation+DP、Diffusion+DP 四类方法。",
        ],
        [
            "设置 epsilon=0.1、0.5、1.0",
            "已完成。所有方法、数据集、下游模型均覆盖三档隐私预算。",
        ],
        [
            "随机选取 5 张合成图并与欧式距离最近原图对比",
            "已完成。每组图保存在 visuals 子目录，左侧为 real-nearest，右侧为 synthetic。",
        ],
        [
            "记录生成时间",
            "已完成。字段为 synthesis_seconds。",
        ],
        [
            "记录不同图片数量下生成时间",
            "已完成。k=160 与 k=320 均有记录，可比较合成样本规模对耗时影响。",
        ],
        [
            "记录 CNN/ResNet 训练时间",
            "已完成。字段为 train_seconds，classifier 字段区分 cnn 与 resnet20。",
        ],
        [
            "记录 CNN/ResNet 预测精度",
            "已完成。original_acc 为原始验证集精度，synthetic_acc 为合成验证集精度。",
        ],
    ]
    add_table(doc, rows, ["任务要求", "完成情况"])

    doc.add_heading("2. 实验设计与过程", level=1)
    doc.add_heading("2.1 场景定义", level=2)
    p(
        doc,
        "根据实验任务和视频记录，本研究采用“本地脱敏、远端训练”的工程路径。"
        "企业侧拥有原始图像数据，但不直接上传原始数据；本地使用差分隐私噪声、生成式模型或数据集蒸馏生成可替代训练集，"
        "再将合成训练集交给第三方算力训练 CNN 或 ResNet20。该设计对应视频记录中对同态加密、TEE 与 DP 路径的比较结论："
        "DP 合成数据路线牺牲一定精度，但在工程开销与隐私保护之间更易落地。",
    )

    doc.add_heading("2.2 数据集与预处理", level=2)
    p(
        doc,
        "MNIST 被调整为 3x32x32 张量以复用统一 CNN/ResNet20 管线；CIFAR10 使用原始 32x32 RGB 图像并进行标准归一化。"
        "CIFAR10 的读取过程中发现系统拒绝访问解压目录，因此新增 tar 包直接读取逻辑，从 cifar-10-python.tar.gz 直接加载 batch 到内存，"
        "避免落地解压目录造成的权限中断。",
    )

    doc.add_heading("2.3 脱敏与合成方法", level=2)
    rows = [
        ["DP", "按类别采样真实图片，并注入与 epsilon 相关的高斯噪声。", "生成极快，作为低成本基线。"],
        ["GAN+DP", "训练轻量生成器/判别器，并对梯度注入 DP 噪声。", "本轮表现接近随机水平，说明轻量 GAN 不适合当前设置。"],
        ["Dataset Distillation+DP", "按类构造原型图片并加噪。", "合成域精度高，但真实域泛化不稳定。"],
        ["Diffusion+DP", "使用逐步噪声-恢复式近似扩散合成，并注入 DP 噪声。", "两数据集平均精度最高。"],
    ]
    add_table(doc, rows, ["方法", "实现路径", "实验观察"])

    doc.add_heading("2.4 训练与评估协议", level=2)
    p(
        doc,
        "下游模型保持原始 CNN 与 ResNet20 不变。每组实验按 method、epsilon、k、classifier 组成一个训练任务。"
        "合成数据按类别拆分为 80% 训练集与 20% 合成验证集，避免在同一批合成训练样本上自测导致 synthetic_acc 虚高。"
        "主要评价指标为 original_acc，即模型在原始验证集上的准确率；synthetic_acc 仅作为合成域拟合程度的辅助指标。",
    )
    rows = [
        ["数据集", "MNIST, CIFAR10"],
        ["方法", "DP, GAN+DP, Dataset Distillation+DP, Diffusion+DP"],
        ["epsilon", "0.1, 0.5, 1.0"],
        ["k", "160, 320 images per class"],
        ["下游模型", "cnn, resnet20"],
        ["训练轮数", "5 epochs"],
        ["合成验证比例", "20%"],
        ["主要指标", "original_acc"],
        ["辅助指标", "synthetic_acc, synthesis_seconds, train_seconds, prediction seconds, visual comparisons"],
    ]
    add_table(doc, rows, ["设置项", "取值"])

    doc.add_heading("3. 实验结果", level=1)
    doc.add_picture(str(figs["best"]), width=Inches(5.8))
    add_caption(doc, "图 1. MNIST 与 CIFAR10 的最优配置及原始验证集精度。")

    best_rows = []
    for dataset in ["MNIST", "CIFAR10"]:
        row = data[data["dataset"] == dataset].sort_values("original_acc", ascending=False).iloc[0]
        best_rows.append(
            [
                dataset,
                METHOD_LABELS[row["method"]],
                row["epsilon"],
                int(row["per_class"]),
                row["classifier"],
                pct(row["original_acc"]),
                pct(row["synthetic_acc"]),
                sec(row["synthesis_seconds"]),
                sec(row["train_seconds"]),
            ]
        )
    add_table(
        doc,
        best_rows,
        [
            "Dataset",
            "Best method",
            "epsilon",
            "k",
            "classifier",
            "original_acc",
            "synthetic_acc",
            "synthesis(s)",
            "train(s)",
        ],
    )

    top_rows = []
    for dataset in ["MNIST", "CIFAR10"]:
        sub = data[data["dataset"] == dataset].sort_values("original_acc", ascending=False).head(5)
        for _, row in sub.iterrows():
            top_rows.append(
                [
                    dataset,
                    METHOD_LABELS[row["method"]],
                    row["epsilon"],
                    int(row["per_class"]),
                    row["classifier"],
                    pct(row["original_acc"]),
                    pct(row["synthetic_acc"]),
                ]
            )
    add_table(
        doc,
        top_rows,
        ["Dataset", "method", "epsilon", "k", "classifier", "original_acc", "synthetic_acc"],
    )

    doc.add_picture(str(figs["method"]), width=Inches(6.2))
    add_caption(doc, "图 2. 四类脱敏方法在两个数据集上的平均 original_acc。")

    doc.add_picture(str(figs["epsilon"]), width=Inches(6.2))
    add_caption(doc, "图 3. 不同隐私预算 epsilon 下的平均效用变化。")

    doc.add_picture(str(figs["k"]), width=Inches(6.2))
    add_caption(doc, "图 4. 每类合成样本数 k 对 original_acc 的影响。")

    doc.add_picture(str(figs["time"]), width=Inches(6.2))
    add_caption(doc, "图 5. 不同方法的平均生成时间与训练时间比较。")

    doc.add_heading("4. 视觉差异对比", level=1)
    p(
        doc,
        "按照任务要求，每组实验随机选取 5 张合成图，并从原始训练池中寻找欧式距离最近的真实图片进行并排展示。"
        "下图选择 MNIST 与 CIFAR10 的最优 Diffusion+DP 配置作为代表。完整视觉结果位于各数据集 outputs/visuals 子目录。",
    )
    visual_paths = [
        ROOT / "mnist/visuals/mnist/diffusion_dp_eps_0.5_k_320.png",
        ROOT / "cifar10/visuals/cifar10/diffusion_dp_eps_1.0_k_320.png",
    ]
    for idx, img_path in enumerate(visual_paths, 1):
        if img_path.exists():
            doc.add_picture(str(img_path), width=Inches(4.7))
            add_caption(doc, f"图 6.{idx}. 最优 Diffusion+DP 配置的 real-nearest 与 synthetic 视觉对比。")

    doc.add_heading("5. 讨论", level=1)
    doc.add_heading("5.1 方法排序", level=2)
    p(
        doc,
        "在两个数据集上，Diffusion+DP 均取得最高平均 original_acc。MNIST 上 Diffusion+DP 的平均 original_acc 为 91.82%，"
        "显著高于 DP 的 80.67%；CIFAR10 上 Diffusion+DP 的平均 original_acc 为 40.66%，也高于 DP 的 31.52%。"
        "这说明扩散式合成在当前实验管线中更能保留可学习结构。GAN+DP 在当前轻量 GAN 设置下接近随机水平，"
        "不宜作为后续主攻方向。Dataset Distillation+DP 的 synthetic_acc 常接近 100%，但 original_acc 明显较低，表明存在合成域过拟合。",
    )

    doc.add_heading("5.2 隐私预算与样本规模", level=2)
    p(
        doc,
        "epsilon 提高通常降低噪声强度，但精度并不总是单调提升。MNIST 的最优点出现在 epsilon=0.5，"
        "CIFAR10 的最优点出现在 epsilon=1.0。相比之下，扩大 k 的收益更稳定：两个数据集上 k=320 的平均精度均高于 k=160。"
        "因此下一步提高本实验精度时，应优先扩大合成样本规模，再围绕当前最优 epsilon 做局部精扫。",
    )

    doc.add_heading("5.3 数据集难度差异", level=2)
    p(
        doc,
        "MNIST 结构简单，当前管线已经达到 96.91% 的最高精度；CIFAR10 语义和纹理复杂度更高，最高精度为 47.66%。"
        "这说明当前轻量合成器已能支撑 MNIST 级别任务，但 CIFAR10 仍受限于合成样本质量和样本规模。"
        "在不改变下游模型结构的约束下，CIFAR10 后续应优先尝试 k=640/1000、epochs=10/20、epsilon=0.7/1.0/1.5/2.0 的组合。",
    )

    doc.add_heading("6. 结论", level=1)
    p(
        doc,
        "本文完整实现并评估了任务图所要求的本地 DP 脱敏、第三方训练流程。实验结果表明，Diffusion+DP 是当前最优方法，"
        "在 MNIST 上达到 96.91% 原始验证精度，在 CIFAR10 上达到 47.66% 原始验证精度。"
        "扩大每类合成样本数 k 是当前最稳定的提精度手段；epsilon 需要数据集相关的局部调参，而不是简单越大越好。"
        "合成验证集精度必须与原始验证集精度联合分析，尤其是 Dataset Distillation+DP 中 synthetic_acc 接近 100% 时，"
        "若 original_acc 较低，则代表模型仅拟合了合成域而非真实数据分布。"
    )

    doc.add_heading("7. 下一步实验建议", level=1)
    rows = [
        ["优先级 1", "CIFAR10 Diffusion+DP", "epsilon=0.7/1.0/1.5/2.0, k=640/1000, epochs=10", "提升 CIFAR10 original_acc"],
        ["优先级 2", "MNIST Diffusion+DP", "epsilon=0.3/0.5/0.7, k=640, epochs=10", "验证是否能超过 97%"],
        ["优先级 3", "重复实验", "对最优 3 组做 3 seeds", "报告均值与标准差"],
        ["优先级 4", "GAN+DP 降级", "暂不扩展", "当前收益低且耗时高"],
    ]
    add_table(doc, rows, ["优先级", "对象", "配置", "目的"])

    doc.add_heading("附录 A. 结果文件", level=1)
    rows = [
        ["MNIST 指标", str(ROOT / "mnist/metrics.csv")],
        ["CIFAR10 指标", str(ROOT / "cifar10/metrics.csv")],
        ["MNIST 视觉图", str(ROOT / "mnist/visuals/mnist")],
        ["CIFAR10 视觉图", str(ROOT / "cifar10/visuals/cifar10")],
        ["本文档图表", str(FIG_DIR)],
    ]
    add_table(doc, rows, ["内容", "路径"])

    doc.save(DOCX_PATH)


def main() -> None:
    ensure_dirs()
    setup_plot_style()
    data = load_data()
    figs = make_figures(data)
    build_doc(data, figs)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
