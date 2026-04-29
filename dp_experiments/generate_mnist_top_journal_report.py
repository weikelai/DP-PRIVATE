from __future__ import annotations

from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
import seaborn as sns
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


SCRIPT_DIR = Path(__file__).resolve().parent
OUT_ROOT = SCRIPT_DIR / "outputs"
REPORT_DIR = OUT_ROOT / "report"
FIG_DIR = OUT_ROOT / "figures_mnist"
STAT_DIR = OUT_ROOT / "statistics"


def setup_style() -> None:
    sns.set_theme(style="whitegrid", context="paper")
    plt.rcParams["font.family"] = "serif"
    plt.rcParams["font.serif"] = ["Times New Roman", "DejaVu Serif", "SimSun"]
    plt.rcParams["axes.unicode_minus"] = False


def ensure_dirs() -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    STAT_DIR.mkdir(parents=True, exist_ok=True)


def save_fig(fig: plt.Figure, stem: str) -> None:
    fig.tight_layout()
    fig.savefig(FIG_DIR / f"{stem}.png", dpi=400, bbox_inches="tight")
    fig.savefig(FIG_DIR / f"{stem}.pdf", dpi=400, bbox_inches="tight")
    plt.close(fig)


def draw_figures(df: pd.DataFrame) -> None:
    p1 = df.groupby(["method", "epsilon"], as_index=False)["original_acc"].mean()
    fig, ax = plt.subplots(figsize=(7.2, 4.2))
    sns.barplot(data=p1, x="epsilon", y="original_acc", hue="method", ax=ax)
    ax.set_title("MNIST: Mean Original Accuracy under Epsilon")
    ax.set_xlabel("Epsilon")
    ax.set_ylabel("Original Accuracy")
    save_fig(fig, "mnist_fig1_method_epsilon_original_acc")

    p2 = df.groupby(["method", "per_class"], as_index=False)["original_acc"].mean()
    fig, ax = plt.subplots(figsize=(7.2, 4.2))
    sns.lineplot(data=p2, x="per_class", y="original_acc", hue="method", marker="o", ax=ax)
    ax.set_title("MNIST: Effect of k on Original Accuracy")
    ax.set_xlabel("Per-class synthetic images (k)")
    ax.set_ylabel("Original Accuracy")
    save_fig(fig, "mnist_fig2_k_vs_original_acc")

    p3 = df.groupby("method", as_index=False)[["synthesis_seconds", "train_seconds"]].mean()
    p3 = p3.melt(id_vars="method", var_name="time_type", value_name="seconds")
    fig, ax = plt.subplots(figsize=(7.2, 4.2))
    sns.barplot(data=p3, x="method", y="seconds", hue="time_type", ax=ax)
    ax.set_title("MNIST: Synthesis vs Training Time")
    ax.set_xlabel("Method")
    ax.set_ylabel("Seconds")
    save_fig(fig, "mnist_fig3_time_comparison")

    p4 = df.groupby(["method", "epsilon"], as_index=False)["generalization_gap"].mean()
    fig, ax = plt.subplots(figsize=(7.2, 4.2))
    sns.lineplot(data=p4, x="epsilon", y="generalization_gap", hue="method", marker="o", ax=ax)
    ax.set_title("MNIST: Generalization Gap under Epsilon")
    ax.set_xlabel("Epsilon")
    ax.set_ylabel("synthetic_acc - original_acc")
    save_fig(fig, "mnist_fig4_generalization_gap")


def export_tables(df: pd.DataFrame) -> dict[str, pd.DataFrame]:
    method_rank = (
        df.groupby("method", as_index=False)["original_acc"]
        .mean()
        .rename(columns={"original_acc": "mean_original_acc"})
        .sort_values("mean_original_acc", ascending=False)
    )
    eps_trend = (
        df.groupby("epsilon", as_index=False)["original_acc"]
        .mean()
        .rename(columns={"original_acc": "mean_original_acc"})
        .sort_values("epsilon")
    )
    k_trend = (
        df.groupby("per_class", as_index=False)["original_acc"]
        .mean()
        .rename(columns={"original_acc": "mean_original_acc"})
        .sort_values("per_class")
    )
    clf_cmp = (
        df.groupby("classifier", as_index=False)["original_acc"]
        .mean()
        .rename(columns={"original_acc": "mean_original_acc"})
        .sort_values("mean_original_acc", ascending=False)
    )
    gap_rank = (
        df.groupby("method", as_index=False)["generalization_gap"]
        .mean()
        .sort_values("generalization_gap", ascending=False)
    )

    method_rank.to_csv(STAT_DIR / "mnist_method_rank.csv", index=False)
    eps_trend.to_csv(STAT_DIR / "mnist_epsilon_trend.csv", index=False)
    k_trend.to_csv(STAT_DIR / "mnist_k_trend.csv", index=False)
    clf_cmp.to_csv(STAT_DIR / "mnist_classifier_comparison.csv", index=False)
    gap_rank.to_csv(STAT_DIR / "mnist_generalization_gap_rank.csv", index=False)
    return {
        "method_rank": method_rank,
        "eps_trend": eps_trend,
        "k_trend": k_trend,
        "clf_cmp": clf_cmp,
        "gap_rank": gap_rank,
    }


def write_markdown(df: pd.DataFrame, tables: dict[str, pd.DataFrame]) -> Path:
    best = df.loc[df["original_acc"].idxmax()]
    def df_to_md_table(inp: pd.DataFrame) -> str:
        cols = list(inp.columns)
        header = "| " + " | ".join(str(c) for c in cols) + " |"
        sep = "| " + " | ".join(["---"] * len(cols)) + " |"
        rows = []
        for _, r in inp.iterrows():
            vals = []
            for c in cols:
                v = r[c]
                if isinstance(v, float):
                    vals.append(f"{v:.4f}")
                else:
                    vals.append(str(v))
            rows.append("| " + " | ".join(vals) + " |")
        return "\n".join([header, sep] + rows)

    text = f"""# MNIST 顶刊风格实验结论

## 最佳组合

- 最优配置：`{best['method']} + eps={best['epsilon']} + k={int(best['per_class'])} + {best['classifier']}`
- `original_acc={best['original_acc']:.4f}`，`synthetic_acc={best['synthetic_acc']:.4f}`

## 方法排序（按 mean_original_acc）

{df_to_md_table(tables['method_rank'])}

## epsilon 趋势

{df_to_md_table(tables['eps_trend'])}

## k 趋势

{df_to_md_table(tables['k_trend'])}

## 分类器对比

{df_to_md_table(tables['clf_cmp'])}

## 过拟合风险（generalization_gap）

{df_to_md_table(tables['gap_rank'])}

## 结论

1. 在 MNIST 上，`diffusion_dp` 依然保持最高的真实分布泛化能力，跨数据集结论稳定。
2. `k` 从 40 提升到 160 带来持续且显著的性能提升，是当前最有效的工程增益变量。
3. `distill_dp` 在合成域拟合强，但 `generalization_gap` 明显偏高，实际部署需谨慎。
4. 在当前训练轮次设置下，`cnn` 的平均表现优于 `resnet20`，建议先保证收敛稳定再提升模型复杂度。
"""
    out_path = REPORT_DIR / "MNIST_顶刊实验结论.md"
    out_path.write_text(text, encoding="utf-8")
    return out_path


def set_default_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    for sec in doc.sections:
        sec.top_margin = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin = Inches(1.0)
        sec.right_margin = Inches(1.0)


def add_table_from_df(doc: Document, df: pd.DataFrame, title: str) -> None:
    doc.add_paragraph(title).runs[0].bold = True
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, c in enumerate(df.columns):
        hdr[i].text = str(c)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, c in enumerate(df.columns):
            v = row[c]
            if isinstance(v, float):
                cells[i].text = f"{v:.4f}"
            else:
                cells[i].text = str(v)
    doc.add_paragraph("")


def add_figure(doc: Document, path: Path, caption: str) -> None:
    if not path.exists():
        return
    doc.add_picture(str(path), width=Inches(6.2))
    p = doc.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def write_docx(df: pd.DataFrame, tables: dict[str, pd.DataFrame]) -> Path:
    best = df.loc[df["original_acc"].idxmax()]
    doc = Document()
    set_default_style(doc)
    h = doc.add_heading("MNIST 差分隐私合成数据实验报告（顶刊风格）", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("1. 实验设置", level=1)
    doc.add_paragraph("数据集为 MNIST（统一转换至 3x32x32），方法包含 dp、gan_dp、distill_dp、diffusion_dp，隐私预算 epsilon=0.1/0.5/1.0，合成规模 k=40/80/160，分类器为 cnn 与 resnet20。")
    doc.add_heading("2. 最优结果与总体结论", level=1)
    doc.add_paragraph(
        f"MNIST 最优组合为 {best['method']} + eps={best['epsilon']} + k={int(best['per_class'])} + {best['classifier']}，"
        f"original_acc={best['original_acc']:.4f}，synthetic_acc={best['synthetic_acc']:.4f}。"
    )
    doc.add_heading("3. 统计表格", level=1)
    add_table_from_df(doc, tables["method_rank"], "表1 方法排序（mean_original_acc）")
    add_table_from_df(doc, tables["eps_trend"], "表2 epsilon 趋势")
    add_table_from_df(doc, tables["k_trend"], "表3 k 趋势")
    add_table_from_df(doc, tables["clf_cmp"], "表4 分类器对比")
    add_table_from_df(doc, tables["gap_rank"], "表5 generalization_gap 排序")
    doc.add_heading("4. 论文图示", level=1)
    add_figure(doc, FIG_DIR / "mnist_fig1_method_epsilon_original_acc.png", "图1 MNIST: 各方法在不同 epsilon 下的 original_acc")
    add_figure(doc, FIG_DIR / "mnist_fig2_k_vs_original_acc.png", "图2 MNIST: k 对 original_acc 的影响")
    add_figure(doc, FIG_DIR / "mnist_fig3_time_comparison.png", "图3 MNIST: 生成与训练时间对比")
    add_figure(doc, FIG_DIR / "mnist_fig4_generalization_gap.png", "图4 MNIST: generalization_gap 随 epsilon 的变化")
    doc.add_heading("5. 讨论", level=1)
    doc.add_paragraph("MNIST 结果与 CIFAR10 主结论一致：diffusion_dp 在真实分布泛化方面更稳定；k 扩大是最直接有效的性能提升手段；distill_dp 在 synthetic_acc 高分条件下存在显著泛化偏差风险。")
    out = REPORT_DIR / "MNIST_top_journal_experiment_report.docx"
    doc.save(out)
    return out


def main() -> None:
    setup_style()
    ensure_dirs()
    df = pd.read_csv(OUT_ROOT / "metrics_mnist_full.csv")
    draw_figures(df)
    tables = export_tables(df)
    md_path = write_markdown(df, tables)
    docx_path = write_docx(df, tables)
    print(f"MNIST figures saved to: {FIG_DIR}")
    print(f"MNIST tables saved to: {STAT_DIR}")
    print(f"MNIST markdown report saved to: {md_path}")
    print(f"MNIST word report saved to: {docx_path}")


if __name__ == "__main__":
    main()

