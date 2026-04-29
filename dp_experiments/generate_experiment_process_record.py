from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Iterable

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path("dp_experiments")
OUT = ROOT / "outputs"
REPORT_DIR = OUT / "next_priority_experiments" / "report"
DOCX_PATH = REPORT_DIR / "DiffusionDP_完整实验流程记录.docx"
DOCX_FALLBACK_PATH = REPORT_DIR / "DiffusionDP_完整实验流程记录_v2.docx"
MD_PATH = REPORT_DIR / "DiffusionDP_完整实验流程记录.md"


def pct(value: float) -> str:
    return f"{float(value) * 100:.2f}%"


def pp(value: float) -> str:
    return f"{float(value) * 100:.2f} pp"


def seconds(value: float) -> str:
    return f"{float(value):.2f}s"


def read_csv(path: str | Path) -> pd.DataFrame:
    return pd.read_csv(path)


def add_table(doc: Document, headers: list[str], rows: Iterable[Iterable[object]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = str(header)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    doc.add_paragraph("")


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def df_best_rows() -> dict[str, pd.Series]:
    complete_mnist = read_csv(OUT / "complete_original_models" / "mnist" / "metrics.csv")
    complete_cifar = read_csv(OUT / "complete_original_models" / "cifar10" / "metrics.csv")
    next_all = read_csv(OUT / "next_priority_experiments" / "diffusion_sweep_all.csv")
    best = {
        "complete_mnist": complete_mnist.loc[complete_mnist["original_acc"].idxmax()],
        "complete_cifar10": complete_cifar.loc[complete_cifar["original_acc"].idxmax()],
        "next_cifar10": next_all[next_all["dataset"] == "cifar10"].loc[
            next_all[next_all["dataset"] == "cifar10"]["original_acc"].idxmax()
        ],
        "next_mnist": next_all[next_all["dataset"] == "mnist"].loc[
            next_all[next_all["dataset"] == "mnist"]["original_acc"].idxmax()
        ],
    }
    return best


def build_markdown() -> str:
    best = df_best_rows()
    repeats = read_csv(OUT / "next_priority_experiments" / "repeats_summary.csv")
    complete_mnist = read_csv(OUT / "complete_original_models" / "mnist" / "metrics.csv")
    complete_cifar = read_csv(OUT / "complete_original_models" / "cifar10" / "metrics.csv")
    next_all = read_csv(OUT / "next_priority_experiments" / "diffusion_sweep_all.csv")
    next_balance = read_csv(OUT / "next_balance_sweep" / "metrics.csv")
    old_repeats = read_csv(OUT / "metrics_resnet20_top2_repeats_summary.csv")
    quality = read_csv(OUT / "metrics_image_quality_corrected.csv")

    lines: list[str] = []
    lines.append("# Diffusion+DP 完整实验流程记录")
    lines.append("")
    lines.append(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append("")
    lines.append("## 1. 记录范围")
    lines.append("")
    lines.append(
        "本文档整合 `dp_experiments` 目录下已有进度记录、阶段性 Markdown、Word 报告、CSV 指标表和最新 "
        "`DiffusionDP_next_priority_report.docx`，形成从实验骨架、全量方法比较、评估口径修正、完整模型矩阵到最终 "
        "Diffusion+DP 优先实验的流程记录。"
    )
    lines.append("")
    lines.append("## 2. 核心结论")
    lines.append("")
    lines.append(
        f"- 完整四方法矩阵阶段：MNIST 最优为 Diffusion+DP, eps={best['complete_mnist'].epsilon:g}, "
        f"k={int(best['complete_mnist'].per_class)}, {best['complete_mnist'].classifier}, "
        f"original_acc={pct(best['complete_mnist'].original_acc)}；CIFAR10 最优为 Diffusion+DP, "
        f"eps={best['complete_cifar10'].epsilon:g}, k={int(best['complete_cifar10'].per_class)}, "
        f"{best['complete_cifar10'].classifier}, original_acc={pct(best['complete_cifar10'].original_acc)}。"
    )
    lines.append(
        f"- 最新 Diffusion+DP 优先实验：CIFAR10 单次最优提升到 {pct(best['next_cifar10'].original_acc)}，"
        f"MNIST 单次最优提升到 {pct(best['next_mnist'].original_acc)}。"
    )
    lines.append(
        "- 3 seeds 重复实验显示：CIFAR10 eps=1.5,k=1000,CNN 均值为 "
        f"{pct(repeats.iloc[0].original_acc_mean)} ± {pp(repeats.iloc[0].original_acc_std)}；"
        "MNIST eps=0.5,k=640,ResNet20 均值为 "
        f"{pct(repeats[repeats.dataset == 'mnist'].iloc[0].original_acc_mean)} ± "
        f"{pp(repeats[repeats.dataset == 'mnist'].iloc[0].original_acc_std)}。"
    )
    lines.append("")
    lines.append("## 3. 源材料清单")
    lines.append("")
    source_rows = [
        ("项目说明", "dp_experiments/README.md", "实验骨架、目录和快速运行说明"),
        ("进度记录", "dp_experiments/EXPERIMENT_PROGRESS.md", "早期全量矩阵、重复实验、论文图表自动化记录"),
        ("阶段修正", "outputs/next_balance_sweep/实验修改说明与下一步方案.md", "synthetic_acc 拆分、视觉图统计修正、下一步调参建议"),
        ("完整矩阵报告", "outputs/complete_original_models/report/DP_synthetic_training_top_journal_report.docx", "MNIST/CIFAR10 完整四方法矩阵报告"),
        ("最新报告", "outputs/next_priority_experiments/report/DiffusionDP_next_priority_report.docx", "Diffusion+DP 最终优先实验结果"),
        ("最新汇总", "outputs/next_priority_experiments/下一步实验结果汇总.md", "最新 sweep 和 3 seeds 结果摘要"),
    ]
    lines.append("| 类型 | 路径 | 用途 |")
    lines.append("|---|---|---|")
    for row in source_rows:
        lines.append("| " + " | ".join(row) + " |")
    lines.append("")
    lines.append("## 4. 统一实验流程")
    lines.append("")
    flow = [
        "加载数据：MNIST 转为 3x32x32；CIFAR10 使用 32x32 RGB；训练集按 train_ratio=0.8 划分本地训练池与原始验证集。",
        "本地脱敏/合成：在企业侧对训练池执行 DP、GAN+DP、Dataset Distillation+DP、Diffusion+DP 中的一种方法，生成每类 k 张合成图。",
        "视觉核查：每组随机选取 5 张合成图，并在本地原始训练池中用欧氏距离寻找最近邻，保存 real-nearest vs synthetic 对比图。",
        "合成域拆分：将合成数据按类别拆成 80% 训练集与 20% 合成验证集，避免 synthetic_acc 在同一训练集上自测而虚高。",
        "第三方训练：只把合成训练集交给 CNN、ResNet18 或 ResNet20 训练；主线实际结果以 CNN 与 ResNet20 为主。",
        "双域评估：训练后分别在原始验证集和合成验证集上评估 original_acc 与 synthetic_acc，并记录预测耗时。",
        "结果归档：每组写入 metrics.csv，重复实验写入 repeats_detail.csv/repeats_summary.csv，报告脚本生成 Markdown、Word 与图表。",
    ]
    for i, item in enumerate(flow, 1):
        lines.append(f"{i}. {item}")
    lines.append("")
    lines.append("## 5. 阶段一：可运行骨架与初始全量矩阵")
    lines.append("")
    lines.append(
        "初始目标是按任务图覆盖四类脱敏方法、三档 epsilon、不同 k、CNN/ResNet 系列和时间/精度/视觉指标。"
        "早期主表 `metrics_task_full.csv` 共 72 组，覆盖 method={dp, gan_dp, distill_dp, diffusion_dp}、"
        "epsilon={0.1,0.5,1.0}、k={40,80,160}、classifier={cnn,resnet20}。"
    )
    first = read_csv(OUT / "metrics_task_full.csv")
    first_best = first.loc[first["original_acc"].idxmax()]
    lines.append(
        f"该阶段最优单点为 {first_best.method}, eps={first_best.epsilon:g}, "
        f"k={int(first_best.per_class)}, {first_best.classifier}, original_acc={pct(first_best.original_acc)}。"
    )
    lines.append("")
    lines.append("## 6. 阶段二：ResNet20 top2 重复实验与统计检验")
    lines.append("")
    old_best = old_repeats.sort_values("original_acc_mean", ascending=False).iloc[0]
    lines.append(
        "在初始矩阵基础上，选择 DP 与 Diffusion+DP 两个候选方法，在 ResNet20 上做 k=40/80/160、"
        "epochs=5/10/20、epsilon=0.1/0.5/1.0 的 3 次重复实验。"
    )
    lines.append(
        f"重复实验最佳为 {old_best.method}, eps={old_best.epsilon:g}, k={int(old_best.per_class)}, "
        f"epochs={int(old_best.epochs)}, original_acc_mean={pct(old_best.original_acc_mean)} ± "
        f"{pp(old_best.original_acc_std)}。"
    )
    lines.append("论文图表脚本还完成了 Diffusion+DP vs DP 的 Welch t-test：27 个分组中 8 个分组显著优于 DP，占 29.63%。")
    lines.append("")
    lines.append("## 7. 阶段三：评估口径修正")
    lines.append("")
    lines.append(
        "检查发现，旧版 synthetic_acc 曾在训练分类器使用的同一批合成样本上评估，容易反映训练集拟合度而不是合成域泛化。"
        "后续修正为 `stratified_split_tensors` 按类别拆分合成训练/验证集，默认 synthetic_eval_ratio=0.2。"
    )
    lines.append(
        "视觉差异统计也从旧的 8 行切分修正为当前保存逻辑中的 5 个随机样本，并裁掉标题和白边后重新计算 SSIM/PSNR。"
    )
    best_quality = quality.sort_values(["epsilon", "ssim"], ascending=[True, False]).groupby("epsilon").head(1)
    lines.append("修正后 SSIM 每档 epsilon 的最高项为：")
    lines.append("| epsilon | method | SSIM | PSNR |")
    lines.append("|---:|---|---:|---:|")
    for _, r in best_quality.iterrows():
        lines.append(f"| {r.epsilon:g} | {r.method} | {r.ssim:.4f} | {r.psnr:.4f} |")
    lines.append("")
    lines.append("## 8. 阶段四：完整四方法矩阵")
    lines.append("")
    lines.append(
        "修正口径后，完整矩阵统一使用 MNIST 与 CIFAR10、四类方法、epsilon=0.1/0.5/1.0、k=160/320、"
        "CNN/ResNet20、epochs=5。该阶段明确 original_acc 是主指标，synthetic_acc 是辅助指标。"
    )
    lines.append("| 数据集 | 最优方法 | epsilon | k | 分类器 | original_acc | synthetic_acc | train_seconds |")
    lines.append("|---|---|---:|---:|---|---:|---:|---:|")
    for name, row in [("MNIST", best["complete_mnist"]), ("CIFAR10", best["complete_cifar10"])]:
        lines.append(
            f"| {name} | {row.method} | {row.epsilon:g} | {int(row.per_class)} | {row.classifier} | "
            f"{pct(row.original_acc)} | {pct(row.synthetic_acc)} | {seconds(row.train_seconds)} |"
        )
    lines.append("")
    lines.append("方法平均 original_acc：")
    lines.append("| 数据集 | method | mean_original_acc |")
    lines.append("|---|---|---:|")
    for dataset, df in [("MNIST", complete_mnist), ("CIFAR10", complete_cifar)]:
        for method, value in df.groupby("method")["original_acc"].mean().sort_values(ascending=False).items():
            lines.append(f"| {dataset} | {method} | {pct(value)} |")
    lines.append("")
    lines.append("## 9. 阶段五：MNIST 平衡小扫")
    lines.append("")
    nb_best = next_balance.loc[next_balance["original_acc"].idxmax()]
    lines.append(
        "在完整矩阵后，先对 MNIST 的 DP 与 Diffusion+DP 做 epsilon=0.3/0.5/0.7、k=160、CNN、epochs=3 的小扫，"
        "验证 epsilon 并非越大越好。"
    )
    lines.append(
        f"该小扫最优为 {nb_best.method}, eps={nb_best.epsilon:g}, k={int(nb_best.per_class)}, "
        f"{nb_best.classifier}, original_acc={pct(nb_best.original_acc)}。"
    )
    lines.append("")
    lines.append("## 10. 阶段六：最新 Diffusion+DP 优先实验")
    lines.append("")
    lines.append(
        "最新报告只扩展 Diffusion+DP，保持下游模型结构不变。CIFAR10 扫描 epsilon=0.7/1.0/1.5/2.0、"
        "k=640/1000、epochs=10；MNIST 扫描 epsilon=0.3/0.5/0.7、k=640、epochs=10。"
    )
    lines.append("| 数据集 | method | epsilon | k | classifier | original_acc | synthetic_acc | train_seconds |")
    lines.append("|---|---|---:|---:|---|---:|---:|---:|")
    for name, row in [("CIFAR10", best["next_cifar10"]), ("MNIST", best["next_mnist"])]:
        lines.append(
            f"| {name} | {row.method} | {row.epsilon:g} | {int(row.per_class)} | {row.classifier} | "
            f"{pct(row.original_acc)} | {pct(row.synthetic_acc)} | {seconds(row.train_seconds)} |"
        )
    lines.append("")
    lines.append("3 seeds 重复实验：")
    lines.append("| 数据集 | config | n | original_acc mean±std | synthetic_acc mean±std | train_seconds_mean |")
    lines.append("|---|---|---:|---:|---:|---:|")
    for _, r in repeats.iterrows():
        lines.append(
            f"| {r.dataset.upper()} | {r.config_id} | {int(r.n)} | {pct(r.original_acc_mean)} ± {pp(r.original_acc_std)} | "
            f"{pct(r.synthetic_acc_mean)} ± {pp(r.synthetic_acc_std)} | {seconds(r.train_seconds_mean)} |"
        )
    lines.append("")
    lines.append("本阶段相对上一轮完整矩阵的提升：CIFAR10 从 47.66% 到 61.40%，提升 +13.74 pp；MNIST 从 96.91% 到 98.72%，提升 +1.82 pp。")
    lines.append("")
    lines.append("## 11. 关键判断")
    lines.append("")
    judgements = [
        "Diffusion+DP 是当前主线最优方法：在早期、完整矩阵和最新优先实验中均给出最高或最稳定的 original_acc。",
        "扩大 k 的收益比盲目增大 epsilon 更稳定。CIFAR10 主要提升来自 k 从 320 扩大到 1000，MNIST 在 k=640 时已接近饱和。",
        "epsilon 存在数据集相关的局部最优。MNIST 多次显示 eps=0.5 较优；CIFAR10 最新重复实验中 eps=1.0 与 1.5 均值接近。",
        "GAN+DP 在当前轻量实现下生成耗时高、精度接近随机，后续不作为优先扩展对象。",
        "Dataset Distillation+DP 可能出现 synthetic_acc 高而 original_acc 低，应始终以 original_acc 为主指标。",
        "ResNet18 是任务图要求的一类 ResNet 系列模型，当前正式重复主线主要是 ResNet20；如写论文或最终汇报，应说明这一点或补跑 ResNet18。",
    ]
    for item in judgements:
        lines.append(f"- {item}")
    lines.append("")
    lines.append("## 12. 后续实验建议")
    lines.append("")
    next_steps = [
        "CIFAR10 小网格：epsilon=1.0/1.25/1.5，k=1000/1500，epochs=12/15，优先确认 60% 左右结果能否继续提升。",
        "在不改模型结构前提下加入训练侧增强：RandomCrop、HorizontalFlip、label_smoothing=0.03/0.05、lr=5e-4/1e-3、weight_decay=5e-5/1e-4。",
        "对 CIFAR10 最优 2-3 个配置继续做 seeds=42/43/44/45/46，报告均值、标准差和置信区间。",
        "补跑 ResNet18 正式矩阵，补齐任务图与报告口径。",
        "若目标转向论文级结果，应引入真实 DP-SGD 扩散训练、更大 U-Net、公共预训练、更多合成样本和更强下游分类器。",
    ]
    for i, item in enumerate(next_steps, 1):
        lines.append(f"{i}. {item}")
    lines.append("")
    lines.append("## 13. 主要输出路径")
    lines.append("")
    paths = [
        "dp_experiments/outputs/complete_original_models/mnist/metrics.csv",
        "dp_experiments/outputs/complete_original_models/cifar10/metrics.csv",
        "dp_experiments/outputs/next_priority_experiments/diffusion_sweep_all.csv",
        "dp_experiments/outputs/next_priority_experiments/repeats_detail.csv",
        "dp_experiments/outputs/next_priority_experiments/repeats_summary.csv",
        "dp_experiments/outputs/next_priority_experiments/report/DiffusionDP_next_priority_report.docx",
        "dp_experiments/outputs/next_priority_experiments/report/DiffusionDP_完整实验流程记录.docx",
    ]
    for path in paths:
        lines.append(f"- `{path}`")
    lines.append("")
    return "\n".join(lines)


def build_docx() -> Path:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    best = df_best_rows()
    repeats = read_csv(OUT / "next_priority_experiments" / "repeats_summary.csv")
    complete_mnist = read_csv(OUT / "complete_original_models" / "mnist" / "metrics.csv")
    complete_cifar = read_csv(OUT / "complete_original_models" / "cifar10" / "metrics.csv")
    old_repeats = read_csv(OUT / "metrics_resnet20_top2_repeats_summary.csv")
    quality = read_csv(OUT / "metrics_image_quality_corrected.csv")

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "SimSun"
    normal.font.size = Pt(10.5)

    title = doc.add_heading("Diffusion+DP 完整实验流程记录", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("1. 记录范围", level=1)
    doc.add_paragraph(
        "本文档整合 dp_experiments 目录下已有进度记录、阶段性 Markdown、Word 报告、CSV 指标表和最新 "
        "DiffusionDP_next_priority_report.docx，形成从实验骨架、全量方法比较、评估口径修正、完整模型矩阵到最终 "
        "Diffusion+DP 优先实验的流程记录。"
    )

    doc.add_heading("2. 核心结论", level=1)
    add_bullets(
        doc,
        [
            f"完整四方法矩阵阶段：MNIST 最优为 Diffusion+DP, eps={best['complete_mnist'].epsilon:g}, "
            f"k={int(best['complete_mnist'].per_class)}, {best['complete_mnist'].classifier}, "
            f"original_acc={pct(best['complete_mnist'].original_acc)}；CIFAR10 最优为 Diffusion+DP, "
            f"eps={best['complete_cifar10'].epsilon:g}, k={int(best['complete_cifar10'].per_class)}, "
            f"{best['complete_cifar10'].classifier}, original_acc={pct(best['complete_cifar10'].original_acc)}。",
            f"最新 Diffusion+DP 优先实验：CIFAR10 单次最优为 {pct(best['next_cifar10'].original_acc)}，"
            f"MNIST 单次最优为 {pct(best['next_mnist'].original_acc)}。",
            "3 seeds 重复实验显示 CIFAR10 eps=1.5,k=1000,CNN 与 eps=1.0,k=1000,CNN 均值接近，"
            "MNIST eps=0.5,k=640,ResNet20 稳定在 98.5% 左右。",
        ],
    )

    doc.add_heading("3. 源材料清单", level=1)
    add_table(
        doc,
        ["类型", "路径", "用途"],
        [
            ("项目说明", "dp_experiments/README.md", "实验骨架、目录和快速运行说明"),
            ("进度记录", "dp_experiments/EXPERIMENT_PROGRESS.md", "早期矩阵、重复实验、论文图表自动化"),
            ("阶段修正", "outputs/next_balance_sweep/实验修改说明与下一步方案.md", "评估口径修正和调参建议"),
            ("完整矩阵报告", "outputs/complete_original_models/report/DP_synthetic_training_top_journal_report.docx", "MNIST/CIFAR10 完整四方法矩阵"),
            ("最新报告", "outputs/next_priority_experiments/report/DiffusionDP_next_priority_report.docx", "Diffusion+DP 最终优先实验"),
            ("最新汇总", "outputs/next_priority_experiments/下一步实验结果汇总.md", "最新 sweep 和 3 seeds 摘要"),
        ],
    )

    doc.add_heading("4. 统一实验流程", level=1)
    add_numbered(
        doc,
        [
            "加载数据：MNIST 转为 3x32x32；CIFAR10 使用 32x32 RGB；训练集按 train_ratio=0.8 划分本地训练池与原始验证集。",
            "本地脱敏/合成：在企业侧对训练池执行 DP、GAN+DP、Dataset Distillation+DP、Diffusion+DP 中的一种方法，生成每类 k 张合成图。",
            "视觉核查：每组随机选取 5 张合成图，并在本地原始训练池中用欧氏距离寻找最近邻，保存 real-nearest vs synthetic 对比图。",
            "合成域拆分：将合成数据按类别拆成 80% 训练集与 20% 合成验证集，避免 synthetic_acc 在同一训练集上自测而虚高。",
            "第三方训练：只把合成训练集交给 CNN、ResNet18 或 ResNet20 训练；主线实际结果以 CNN 与 ResNet20 为主。",
            "双域评估：训练后分别在原始验证集和合成验证集上评估 original_acc 与 synthetic_acc，并记录预测耗时。",
            "结果归档：每组写入 metrics.csv，重复实验写入 repeats_detail.csv/repeats_summary.csv，报告脚本生成 Markdown、Word 与图表。",
        ],
    )

    doc.add_heading("5. 阶段一：可运行骨架与初始全量矩阵", level=1)
    first = read_csv(OUT / "metrics_task_full.csv")
    first_best = first.loc[first["original_acc"].idxmax()]
    doc.add_paragraph(
        "初始阶段覆盖 method={dp, gan_dp, distill_dp, diffusion_dp}、epsilon={0.1,0.5,1.0}、"
        "k={40,80,160}、classifier={cnn,resnet20}，共 72 组。"
    )
    doc.add_paragraph(
        f"该阶段最优单点为 {first_best.method}, eps={first_best.epsilon:g}, "
        f"k={int(first_best.per_class)}, {first_best.classifier}, original_acc={pct(first_best.original_acc)}。"
    )

    doc.add_heading("6. 阶段二：ResNet20 top2 重复实验与统计检验", level=1)
    old_best = old_repeats.sort_values("original_acc_mean", ascending=False).iloc[0]
    doc.add_paragraph(
        "选择 DP 与 Diffusion+DP，在 ResNet20 上做 k=40/80/160、epochs=5/10/20、"
        "epsilon=0.1/0.5/1.0 的 3 次重复实验。"
    )
    doc.add_paragraph(
        f"最佳为 {old_best.method}, eps={old_best.epsilon:g}, k={int(old_best.per_class)}, "
        f"epochs={int(old_best.epochs)}, original_acc_mean={pct(old_best.original_acc_mean)} ± "
        f"{pp(old_best.original_acc_std)}。Welch t-test 中 27 个分组有 8 个分组显著优于 DP，占 29.63%。"
    )

    doc.add_heading("7. 阶段三：评估口径修正", level=1)
    doc.add_paragraph(
        "旧版 synthetic_acc 曾在训练分类器使用的同一批合成样本上评估，容易反映训练集拟合度。"
        "后续修正为按类别拆分 80% 合成训练集与 20% 合成验证集。视觉质量统计也修正为 5 个随机样本，"
        "并裁掉标题和白边后重新计算 SSIM/PSNR。"
    )
    best_quality = quality.sort_values(["epsilon", "ssim"], ascending=[True, False]).groupby("epsilon").head(1)
    add_table(
        doc,
        ["epsilon", "method", "SSIM", "PSNR"],
        [(f"{r.epsilon:g}", r.method, f"{r.ssim:.4f}", f"{r.psnr:.4f}") for _, r in best_quality.iterrows()],
    )

    doc.add_heading("8. 阶段四：完整四方法矩阵", level=1)
    doc.add_paragraph(
        "修正口径后，完整矩阵统一使用 MNIST 与 CIFAR10、四类方法、epsilon=0.1/0.5/1.0、"
        "k=160/320、CNN/ResNet20、epochs=5。"
    )
    add_table(
        doc,
        ["数据集", "最优方法", "epsilon", "k", "分类器", "original_acc", "synthetic_acc", "train_seconds"],
        [
            (
                "MNIST",
                best["complete_mnist"].method,
                f"{best['complete_mnist'].epsilon:g}",
                int(best["complete_mnist"].per_class),
                best["complete_mnist"].classifier,
                pct(best["complete_mnist"].original_acc),
                pct(best["complete_mnist"].synthetic_acc),
                seconds(best["complete_mnist"].train_seconds),
            ),
            (
                "CIFAR10",
                best["complete_cifar10"].method,
                f"{best['complete_cifar10'].epsilon:g}",
                int(best["complete_cifar10"].per_class),
                best["complete_cifar10"].classifier,
                pct(best["complete_cifar10"].original_acc),
                pct(best["complete_cifar10"].synthetic_acc),
                seconds(best["complete_cifar10"].train_seconds),
            ),
        ],
    )
    method_rows = []
    for dataset, df in [("MNIST", complete_mnist), ("CIFAR10", complete_cifar)]:
        for method, value in df.groupby("method")["original_acc"].mean().sort_values(ascending=False).items():
            method_rows.append((dataset, method, pct(value)))
    add_table(doc, ["数据集", "method", "mean_original_acc"], method_rows)

    doc.add_heading("9. 阶段五：MNIST 平衡小扫", level=1)
    next_balance = read_csv(OUT / "next_balance_sweep" / "metrics.csv")
    nb_best = next_balance.loc[next_balance["original_acc"].idxmax()]
    doc.add_paragraph(
        "在完整矩阵后，对 MNIST 的 DP 与 Diffusion+DP 做 epsilon=0.3/0.5/0.7、k=160、CNN、epochs=3 的小扫，"
        "用于验证 epsilon 并非越大越好。"
    )
    doc.add_paragraph(
        f"该小扫最优为 {nb_best.method}, eps={nb_best.epsilon:g}, k={int(nb_best.per_class)}, "
        f"{nb_best.classifier}, original_acc={pct(nb_best.original_acc)}。"
    )

    doc.add_heading("10. 阶段六：最新 Diffusion+DP 优先实验", level=1)
    doc.add_paragraph(
        "最新报告只扩展 Diffusion+DP，保持下游模型结构不变。CIFAR10 扫描 epsilon=0.7/1.0/1.5/2.0、"
        "k=640/1000、epochs=10；MNIST 扫描 epsilon=0.3/0.5/0.7、k=640、epochs=10。"
    )
    add_table(
        doc,
        ["数据集", "method", "epsilon", "k", "classifier", "original_acc", "synthetic_acc", "train_seconds"],
        [
            (
                "CIFAR10",
                best["next_cifar10"].method,
                f"{best['next_cifar10'].epsilon:g}",
                int(best["next_cifar10"].per_class),
                best["next_cifar10"].classifier,
                pct(best["next_cifar10"].original_acc),
                pct(best["next_cifar10"].synthetic_acc),
                seconds(best["next_cifar10"].train_seconds),
            ),
            (
                "MNIST",
                best["next_mnist"].method,
                f"{best['next_mnist'].epsilon:g}",
                int(best["next_mnist"].per_class),
                best["next_mnist"].classifier,
                pct(best["next_mnist"].original_acc),
                pct(best["next_mnist"].synthetic_acc),
                seconds(best["next_mnist"].train_seconds),
            ),
        ],
    )
    add_table(
        doc,
        ["数据集", "config", "n", "original_acc mean±std", "synthetic_acc mean±std", "train_seconds_mean"],
        [
            (
                r.dataset.upper(),
                r.config_id,
                int(r.n),
                f"{pct(r.original_acc_mean)} ± {pp(r.original_acc_std)}",
                f"{pct(r.synthetic_acc_mean)} ± {pp(r.synthetic_acc_std)}",
                seconds(r.train_seconds_mean),
            )
            for _, r in repeats.iterrows()
        ],
    )
    doc.add_paragraph(
        "本阶段相对上一轮完整矩阵的提升：CIFAR10 从 47.66% 到 61.40%，提升 +13.74 pp；"
        "MNIST 从 96.91% 到 98.72%，提升 +1.82 pp。"
    )

    for image in [
        OUT / "next_priority_experiments" / "report" / "figures" / "best_result_comparison.png",
        OUT / "next_priority_experiments" / "report" / "figures" / "repeat_mean_std.png",
    ]:
        if image.exists():
            doc.add_picture(str(image), width=Inches(5.8))

    doc.add_heading("11. 关键判断", level=1)
    add_bullets(
        doc,
        [
            "Diffusion+DP 是当前主线最优方法：在早期、完整矩阵和最新优先实验中均给出最高或最稳定的 original_acc。",
            "扩大 k 的收益比盲目增大 epsilon 更稳定。CIFAR10 主要提升来自 k 从 320 扩大到 1000，MNIST 在 k=640 时已接近饱和。",
            "epsilon 存在数据集相关的局部最优。MNIST 多次显示 eps=0.5 较优；CIFAR10 最新重复实验中 eps=1.0 与 1.5 均值接近。",
            "GAN+DP 在当前轻量实现下生成耗时高、精度接近随机，后续不作为优先扩展对象。",
            "Dataset Distillation+DP 可能出现 synthetic_acc 高而 original_acc 低，应始终以 original_acc 为主指标。",
            "ResNet18 是任务图要求的一类 ResNet 系列模型，当前正式重复主线主要是 ResNet20；如写论文或最终汇报，应说明这一点或补跑 ResNet18。",
        ],
    )

    doc.add_heading("12. 后续实验建议", level=1)
    add_numbered(
        doc,
        [
            "CIFAR10 小网格：epsilon=1.0/1.25/1.5，k=1000/1500，epochs=12/15，优先确认 60% 左右结果能否继续提升。",
            "在不改模型结构前提下加入训练侧增强：RandomCrop、HorizontalFlip、label_smoothing=0.03/0.05、lr=5e-4/1e-3、weight_decay=5e-5/1e-4。",
            "对 CIFAR10 最优 2-3 个配置继续做 seeds=42/43/44/45/46，报告均值、标准差和置信区间。",
            "补跑 ResNet18 正式矩阵，补齐任务图与报告口径。",
            "若目标转向论文级结果，应引入真实 DP-SGD 扩散训练、更大 U-Net、公共预训练、更多合成样本和更强下游分类器。",
        ],
    )

    doc.add_heading("13. 主要输出路径", level=1)
    add_bullets(
        doc,
        [
            "dp_experiments/outputs/complete_original_models/mnist/metrics.csv",
            "dp_experiments/outputs/complete_original_models/cifar10/metrics.csv",
            "dp_experiments/outputs/next_priority_experiments/diffusion_sweep_all.csv",
            "dp_experiments/outputs/next_priority_experiments/repeats_detail.csv",
            "dp_experiments/outputs/next_priority_experiments/repeats_summary.csv",
            "dp_experiments/outputs/next_priority_experiments/report/DiffusionDP_next_priority_report.docx",
            "dp_experiments/outputs/next_priority_experiments/report/DiffusionDP_完整实验流程记录.docx",
        ],
    )

    try:
        doc.save(DOCX_PATH)
        return DOCX_PATH
    except PermissionError:
        doc.save(DOCX_FALLBACK_PATH)
        return DOCX_FALLBACK_PATH


def main() -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    MD_PATH.write_text(build_markdown(), encoding="utf-8-sig")
    docx_path = build_docx()
    print(docx_path)
    print(MD_PATH)


if __name__ == "__main__":
    main()
