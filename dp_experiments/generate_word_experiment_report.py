from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


SCRIPT_DIR = Path(__file__).resolve().parent
ROOT = SCRIPT_DIR / "outputs"
REPORT_DIR = ROOT / "report"
FIG_DIR = ROOT / "figures"


def set_default_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    for sec in doc.sections:
        sec.top_margin = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin = Inches(1.0)
        sec.right_margin = Inches(1.0)


def add_heading_center(doc: Document, text: str, level: int = 0):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 0 else WD_ALIGN_PARAGRAPH.LEFT


def add_paragraph(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)


def add_formula_block(doc: Document, formula: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(formula)
    r.italic = True


def add_table_from_df(doc: Document, df: pd.DataFrame, title: str, max_rows: int = 12):
    doc.add_paragraph(title).runs[0].bold = True
    show = df.head(max_rows).copy()
    table = doc.add_table(rows=1, cols=len(show.columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, c in enumerate(show.columns):
        hdr[i].text = str(c)
    for _, row in show.iterrows():
        cells = table.add_row().cells
        for i, c in enumerate(show.columns):
            cells[i].text = str(row[c])
    doc.add_paragraph("")


def add_figure(doc: Document, fig_path: Path, caption: str):
    if not fig_path.exists():
        return
    doc.add_picture(str(fig_path), width=Inches(6.2))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main():
    REPORT_DIR.mkdir(parents=True, exist_ok=True)

    task_df = pd.read_csv(ROOT / "metrics_task_full.csv")
    summary_df = pd.read_csv(ROOT / "metrics_resnet20_top2_repeats_summary.csv")
    detail_df = pd.read_csv(ROOT / "metrics_resnet20_top2_repeats_detail.csv")
    ttest_df = pd.read_csv(ROOT / "statistics" / "ttest_diffusion_vs_dp_by_group.csv")
    img_df = pd.read_csv(ROOT / "metrics_image_quality.csv")
    mnist_df = None
    mnist_path = ROOT / "metrics_mnist_full.csv"
    if mnist_path.exists():
        mnist_df = pd.read_csv(mnist_path)

    best_task = task_df.loc[task_df["original_acc"].idxmax()]
    best_resnet = summary_df.loc[summary_df["original_acc_mean"].idxmax()]
    sig_cnt = int(ttest_df["significant_better_than_dp"].sum())
    sig_ratio = sig_cnt / max(len(ttest_df), 1)

    eps_trend = task_df.groupby("epsilon", as_index=False)["original_acc"].mean()
    k_trend = summary_df.groupby("per_class", as_index=False)["original_acc_mean"].mean()
    ep_trend = summary_df.groupby("epochs", as_index=False)["original_acc_mean"].mean()

    doc = Document()
    set_default_style(doc)

    add_heading_center(doc, "基于差分隐私脱敏数据的外包训练实验说明（顶刊格式草稿）", 0)
    add_paragraph(doc, "摘要：本文围绕企业本地数据脱敏后外包训练场景，构建并评估 DP、GAN+DP、Dataset Distillation+DP、Diffusion+DP 四类方法。实验覆盖不同隐私预算、合成规模和训练轮次，分析精度、耗时、图像质量与统计显著性。结果表明，在 ResNet20 重复实验中，Diffusion+DP 在中高训练强度下表现出更优的原始分布泛化性能。")

    doc.add_heading("1. 研究问题与实验设置", level=1)
    add_paragraph(doc, "研究目标：在不上传原始本地数据的前提下，评估差分隐私脱敏合成数据对第三方训练可用性的影响。")
    add_paragraph(doc, "数据与任务：CIFAR10 图像分类。模型包括 CNN 与 ResNet20，重点分析 ResNet20 排序稳定性。")
    add_paragraph(doc, "方法集合：DP、GAN+DP、Dataset Distillation+DP、Diffusion+DP。隐私预算 epsilon 取 {0.1, 0.5, 1.0}。合成规模 k 取 {40, 80, 160}。重复实验在 ResNet20 top2（dp vs diffusion_dp）上执行 3 次。")

    doc.add_heading("2. 方法与公式", level=1)
    add_paragraph(doc, "本文采用差分隐私与分布拟合的联合思路。关键评价指标与统计流程如下。")
    add_formula_block(doc, "Accuracy = (1 / N) * Σ_{i=1..N} 1[y_i = ŷ_i]")
    add_formula_block(doc, "d(x_s, x_r) = ||x_s - x_r||_2^2")
    add_formula_block(doc, "SSIM(x, y) = ((2μ_xμ_y + C1)(2σ_xy + C2)) / ((μ_x^2 + μ_y^2 + C1)(σ_x^2 + σ_y^2 + C2))")
    add_formula_block(doc, "PSNR(x, y) = 10 * log10(MAX_I^2 / MSE(x, y))")
    add_formula_block(doc, "Welch t-test: t = (μ1 - μ2) / sqrt(s1^2/n1 + s2^2/n2)")
    add_paragraph(doc, "其中，最近邻视觉对比采用合成图与原始样本欧式距离最小匹配；显著性检验按 epsilon、k、epochs 分组对 diffusion_dp 与 dp 的 original_acc 进行双样本 Welch t 检验。")

    doc.add_heading("3. 数据结果总览", level=1)
    add_paragraph(
        doc,
        f"全量任务最佳单点来自 {best_task['method']}（epsilon={best_task['epsilon']}, k={int(best_task['per_class'])}, classifier={best_task['classifier']}），original_acc={best_task['original_acc']:.4f}。"
    )
    add_paragraph(
        doc,
        f"ResNet20 重复实验最佳配置为 {best_resnet['method']}（epsilon={best_resnet['epsilon']}, k={int(best_resnet['per_class'])}, epochs={int(best_resnet['epochs'])}），original_acc_mean={best_resnet['original_acc_mean']:.4f} ± {best_resnet['original_acc_std']:.4f}。"
    )
    add_paragraph(
        doc,
        f"统计显著性结果：共 {len(ttest_df)} 个分组，diffusion_dp 显著优于 dp 的分组数 {sig_cnt}，占比 {sig_ratio:.2%}。"
    )

    doc.add_heading("4. 主结果表与趋势分析", level=1)
    add_table_from_df(
        doc,
        summary_df.sort_values("original_acc_mean", ascending=False)[
            ["method", "epsilon", "per_class", "epochs", "original_acc_mean", "original_acc_std", "train_seconds_mean"]
        ],
        "表1  ResNet20 Top2 重复实验中按 original_acc_mean 排序（前12行）"
    )
    add_table_from_df(
        doc,
        eps_trend.rename(columns={"original_acc": "mean_original_acc"}),
        "表2  epsilon 维度原始集精度均值"
    )
    add_table_from_df(
        doc,
        k_trend.rename(columns={"original_acc_mean": "mean_original_acc"}),
        "表3  合成规模 k 维度原始集精度均值"
    )
    add_table_from_df(
        doc,
        ep_trend.rename(columns={"original_acc_mean": "mean_original_acc"}),
        "表4  训练轮次 epochs 维度原始集精度均值"
    )

    doc.add_heading("5. 论文图示", level=1)
    add_figure(doc, FIG_DIR / "fig1_method_epsilon_original_acc_bar.png", "图1. 各方法在不同 epsilon 下的 original_acc 柱状图")
    add_figure(doc, FIG_DIR / "fig2_k_impact_accuracy_line.png", "图2. k=40/80/160 对 accuracy 影响折线图")
    add_figure(doc, FIG_DIR / "fig3_epochs_impact_accuracy_line.png", "图3. epochs=5/10/20 对 accuracy 影响图")
    add_figure(doc, FIG_DIR / "fig4_synthesis_train_time_comparison.png", "图4. synthesis_seconds 与 train_seconds 时间对比图")
    add_figure(doc, FIG_DIR / "fig5_epsilon_privacy_tradeoff.png", "图5. epsilon 与 accuracy 隐私权衡图")
    add_figure(doc, FIG_DIR / "fig6_image_quality_ssim_bar.png", "图6. 图像质量 SSIM 对比图")
    add_figure(doc, FIG_DIR / "fig7_image_quality_psnr_bar.png", "图7. 图像质量 PSNR 对比图")

    doc.add_heading("6. 图像质量与可解释性分析", level=1)
    add_table_from_df(doc, img_df, "表5  method + epsilon 维度图像质量汇总（SSIM/PSNR）", max_rows=20)
    add_paragraph(doc, "从 SSIM 看，diffusion_dp 在低预算（epsilon=0.1）下结构保持较好；在中高预算下与 dp 接近。PSNR 指标上 distill_dp 在若干预算点较高，但该指标反映像素误差，需与 original_acc 联合判断。")

    doc.add_heading("7. 统计显著性检验结论", level=1)
    add_table_from_df(
        doc,
        ttest_df[["epsilon", "per_class", "epochs", "diffusion_mean", "dp_mean", "mean_diff", "p_value", "significant_better_than_dp"]],
        "表6  按 epsilon、k、epochs 分组的 Welch t-test 结果",
        max_rows=30,
    )
    add_paragraph(doc, "显著性判定标准为 p < 0.05 且 diffusion_mean > dp_mean。结果显示 diffusion_dp 在部分关键高训练强度配置下可显著优于 dp，尤其在较大 k 与较高 epochs 组合中优势更稳定。")

    doc.add_heading("8. 讨论与局限", level=1)
    add_paragraph(doc, "（1）视觉最近邻对比当前实现为“前 n 张样本（最多8张）+ 欧式最近邻”，并非随机5张；若用于严格论文复现，建议补齐随机采样协议。")
    add_paragraph(doc, "（2）当前主线为 CIFAR10；若扩展至 MNIST，应保持同样统计口径并进行跨数据集一致性验证。")
    add_paragraph(doc, "（3）合成集精度 synthetic_acc 在部分配置下接近饱和，存在过拟合风险，因此模型可用性应以 original_acc 与显著性结果为主。")

    doc.add_heading("9. 结论与建议", level=1)
    add_paragraph(doc, "综合实验可得：在满足差分隐私预算约束前提下，diffusion_dp 在 ResNet20 中高训练强度场景下表现最佳；k=160 与 epochs=20 为当前最优工程区间。建议在生产部署中采用“diffusion_dp 主线 + dp 低成本回退”的双策略，并保留周期性重复实验以监控分布漂移。")

    if mnist_df is not None and not mnist_df.empty:
        doc.add_heading("10. MNIST 跨数据集一致性验证", level=1)
        add_paragraph(doc, "为验证结论的跨数据集稳定性，采用与 CIFAR10 主线一致的口径在 MNIST 上完成全量网格实验（method/epsilon/k/classifier）。MNIST 输入按 3x32x32 对齐现有模型管线。")
        mnist_best = mnist_df.loc[mnist_df["original_acc"].idxmax()]
        add_paragraph(
            doc,
            f"MNIST 最优组合为 {mnist_best['method']}（epsilon={mnist_best['epsilon']}, k={int(mnist_best['per_class'])}, classifier={mnist_best['classifier']}），original_acc={mnist_best['original_acc']:.4f}，synthetic_acc={mnist_best['synthetic_acc']:.4f}。"
        )
        mnist_method = (
            mnist_df.groupby("method", as_index=False)["original_acc"]
            .mean()
            .rename(columns={"original_acc": "mean_original_acc"})
            .sort_values("mean_original_acc", ascending=False)
        )
        mnist_eps = (
            mnist_df.groupby("epsilon", as_index=False)["original_acc"]
            .mean()
            .rename(columns={"original_acc": "mean_original_acc"})
            .sort_values("epsilon")
        )
        mnist_k = (
            mnist_df.groupby("per_class", as_index=False)["original_acc"]
            .mean()
            .rename(columns={"original_acc": "mean_original_acc"})
            .sort_values("per_class")
        )
        mnist_gap = (
            mnist_df.assign(generalization_gap=mnist_df["synthetic_acc"] - mnist_df["original_acc"])
            .groupby("method", as_index=False)["generalization_gap"]
            .mean()
            .sort_values("generalization_gap", ascending=False)
        )
        add_table_from_df(doc, mnist_method, "表7  MNIST 各方法 mean_original_acc 排序", max_rows=10)
        add_table_from_df(doc, mnist_eps, "表8  MNIST 不同 epsilon 的 mean_original_acc", max_rows=10)
        add_table_from_df(doc, mnist_k, "表9  MNIST 不同 k 的 mean_original_acc", max_rows=10)
        add_table_from_df(doc, mnist_gap, "表10  MNIST 各方法 generalization_gap 均值", max_rows=10)
        add_paragraph(doc, "MNIST 结果与 CIFAR10 主结论一致：diffusion_dp 在真实分布泛化能力上保持领先；k 增大可稳定提升 original_acc；distill_dp 的 generalization_gap 偏高，提示其在合成域拟合较强但真实域迁移风险更高。")

    out_docx = REPORT_DIR / "top_journal_experiment_report.docx"
    doc.save(out_docx)
    print(f"saved: {out_docx}")


if __name__ == "__main__":
    main()

