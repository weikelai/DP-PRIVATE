from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


SCRIPT_DIR = Path(__file__).resolve().parent
ROOT = SCRIPT_DIR / "outputs"
REPORT_DIR = ROOT / "report"


def set_default_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    for sec in doc.sections:
        sec.top_margin = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin = Inches(1.0)
        sec.right_margin = Inches(1.0)


def add_heading_center(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 0 else WD_ALIGN_PARAGRAPH.LEFT


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)


def add_table_from_df(doc: Document, df: pd.DataFrame, title: str, max_rows: int = 20) -> None:
    doc.add_paragraph(title).runs[0].bold = True
    show = df.head(max_rows).copy()
    table = doc.add_table(rows=1, cols=len(show.columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, c in enumerate(show.columns):
        hdr[i].text = str(c)
    for _, row in show.iterrows():
        cells = table.add_row().cells
        for i, c in enumerate(show.columns):
            v = row[c]
            if isinstance(v, float):
                cells[i].text = f"{v:.4f}"
            else:
                cells[i].text = str(v)
    doc.add_paragraph("")


def build_summary(df: pd.DataFrame, dataset_name: str) -> dict:
    best = df.loc[df["original_acc"].idxmax()]
    method_rank = (
        df.groupby("method", as_index=False)["original_acc"]
        .mean()
        .rename(columns={"original_acc": "mean_original_acc"})
        .sort_values("mean_original_acc", ascending=False)
    )
    eps_trend = (
        df.groupby("epsilon", as_index=False)["original_acc"]
        .mean()
        .rename(columns={"original_acc": "mean_original_acc"})
        .sort_values("epsilon")
    )
    k_trend = (
        df.groupby("per_class", as_index=False)["original_acc"]
        .mean()
        .rename(columns={"original_acc": "mean_original_acc"})
        .sort_values("per_class")
    )
    clf_trend = (
        df.groupby("classifier", as_index=False)["original_acc"]
        .mean()
        .rename(columns={"original_acc": "mean_original_acc"})
        .sort_values("mean_original_acc", ascending=False)
    )
    time_by_method = (
        df.groupby("method", as_index=False)[["synthesis_seconds", "train_seconds"]]
        .mean()
        .sort_values("train_seconds")
    )
    return {
        "dataset": dataset_name,
        "best": best,
        "method_rank": method_rank,
        "eps_trend": eps_trend,
        "k_trend": k_trend,
        "clf_trend": clf_trend,
        "time_by_method": time_by_method,
    }


def main() -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    cifar_df = pd.read_csv(ROOT / "metrics_task_full.csv")
    mnist_df = pd.read_csv(ROOT / "metrics_mnist_full.csv")

    cifar = build_summary(cifar_df, "CIFAR10")
    mnist = build_summary(mnist_df, "MNIST")

    compare_method = (
        cifar["method_rank"]
        .merge(
            mnist["method_rank"],
            on="method",
            suffixes=("_cifar10", "_mnist"),
        )
        .assign(delta_mnist_minus_cifar=lambda d: d["mean_original_acc_mnist"] - d["mean_original_acc_cifar10"])
        .sort_values("mean_original_acc_mnist", ascending=False)
    )
    compare_eps = (
        cifar["eps_trend"]
        .merge(
            mnist["eps_trend"],
            on="epsilon",
            suffixes=("_cifar10", "_mnist"),
        )
        .assign(delta_mnist_minus_cifar=lambda d: d["mean_original_acc_mnist"] - d["mean_original_acc_cifar10"])
    )
    compare_k = (
        cifar["k_trend"]
        .merge(
            mnist["k_trend"],
            on="per_class",
            suffixes=("_cifar10", "_mnist"),
        )
        .assign(delta_mnist_minus_cifar=lambda d: d["mean_original_acc_mnist"] - d["mean_original_acc_cifar10"])
    )

    doc = Document()
    set_default_style(doc)
    add_heading_center(doc, "MNIST 与 CIFAR10 跨数据集对比实验报告", 0)
    add_paragraph(
        doc,
        "摘要：本报告在统一差分隐私实验口径下，对 CIFAR10 与 MNIST 两个数据集的脱敏合成训练结果进行系统对比，重点分析方法排序一致性、隐私预算趋势、样本规模效应、模型差异与时间开销。",
    )

    doc.add_heading("1. 对比设置", level=1)
    add_paragraph(
        doc,
        "统一对比变量：method={dp, gan_dp, distill_dp, diffusion_dp}，epsilon={0.1,0.5,1.0}，k={40,80,160}，classifier={cnn,resnet20}。评估指标为 original_acc、synthetic_acc、generalization_gap 及时间指标。",
    )

    doc.add_heading("2. 最优组合对比", level=1)
    add_paragraph(
        doc,
        f"CIFAR10 最优：{cifar['best']['method']} + eps={cifar['best']['epsilon']} + k={int(cifar['best']['per_class'])} + {cifar['best']['classifier']}，original_acc={cifar['best']['original_acc']:.4f}。",
    )
    add_paragraph(
        doc,
        f"MNIST 最优：{mnist['best']['method']} + eps={mnist['best']['epsilon']} + k={int(mnist['best']['per_class'])} + {mnist['best']['classifier']}，original_acc={mnist['best']['original_acc']:.4f}。",
    )

    doc.add_heading("3. 方法平均精度排序对比", level=1)
    add_table_from_df(doc, compare_method, "表1 各方法 mean_original_acc 跨数据集对比")

    doc.add_heading("4. epsilon 趋势对比", level=1)
    add_table_from_df(doc, compare_eps, "表2 不同 epsilon 的 mean_original_acc 对比")

    doc.add_heading("5. k 趋势对比", level=1)
    add_table_from_df(doc, compare_k, "表3 不同 k 的 mean_original_acc 对比")

    doc.add_heading("6. 单数据集分解结果", level=1)
    add_table_from_df(doc, cifar["clf_trend"], "表4 CIFAR10 分类器对比（mean_original_acc）", max_rows=10)
    add_table_from_df(doc, mnist["clf_trend"], "表5 MNIST 分类器对比（mean_original_acc）", max_rows=10)
    add_table_from_df(doc, cifar["time_by_method"], "表6 CIFAR10 方法时间开销均值（秒）", max_rows=10)
    add_table_from_df(doc, mnist["time_by_method"], "表7 MNIST 方法时间开销均值（秒）", max_rows=10)

    doc.add_heading("7. 结论", level=1)
    add_paragraph(
        doc,
        "两数据集对比显示：diffusion_dp 在真实分布泛化能力方面均保持领先，方法排序具备跨数据集稳定性；k 增大持续带来性能提升，是最稳定有效的工程杠杆；distill_dp 在 synthetic_acc 高分下依然存在 generalization_gap 偏高问题，部署时应以 original_acc 为主并结合显著性验证。",
    )

    out_path = REPORT_DIR / "MNIST_CIFAR10_对比实验报告.docx"
    doc.save(out_path)
    print(f"saved: {out_path}")


if __name__ == "__main__":
    main()

