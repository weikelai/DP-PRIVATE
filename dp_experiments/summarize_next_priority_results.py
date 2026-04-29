from __future__ import annotations

import re
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt


ROOT = Path("dp_experiments/outputs/next_priority_experiments")
PREVIOUS_ROOT = Path("dp_experiments/outputs/complete_original_models")


def pct(value: float) -> str:
    return f"{value * 100:.2f}%"


def load_repeats() -> pd.DataFrame:
    rows: list[pd.DataFrame] = []
    for metrics_path in (ROOT / "repeats").glob("*/metrics.csv"):
        folder = metrics_path.parent.name
        match = re.match(r"(.+)_seed(\d+)$", folder)
        df = pd.read_csv(metrics_path)
        df.insert(0, "config_id", match.group(1) if match else folder)
        df.insert(1, "seed", int(match.group(2)) if match else np.nan)
        df.insert(2, "dataset", "cifar10" if folder.startswith("cifar10") else "mnist")
        df.insert(3, "run_dir", metrics_path.parent.as_posix())
        rows.append(df)
    if not rows:
        raise FileNotFoundError("No repeat metrics.csv files found.")
    return pd.concat(rows, ignore_index=True).sort_values(["dataset", "config_id", "seed"])


def load_sweeps() -> pd.DataFrame:
    rows: list[pd.DataFrame] = []
    for dataset, folder in [("cifar10", "cifar10_diffusion_sweep"), ("mnist", "mnist_diffusion_sweep")]:
        df = pd.read_csv(ROOT / folder / "metrics.csv")
        df.insert(0, "dataset", dataset)
        rows.append(df)
    return pd.concat(rows, ignore_index=True)


def load_previous() -> pd.DataFrame:
    rows: list[pd.DataFrame] = []
    for dataset in ["cifar10", "mnist"]:
        df = pd.read_csv(PREVIOUS_ROOT / dataset / "metrics.csv")
        df.insert(0, "dataset", dataset)
        rows.append(df)
    return pd.concat(rows, ignore_index=True)


def main() -> None:
    ROOT.mkdir(parents=True, exist_ok=True)

    detail = load_repeats()
    detail_path = ROOT / "repeats_detail.csv"
    detail.to_csv(detail_path, index=False, encoding="utf-8-sig")

    summary = (
        detail.groupby(["dataset", "config_id", "method", "epsilon", "per_class", "classifier"])
        .agg(
            n=("original_acc", "count"),
            original_acc_mean=("original_acc", "mean"),
            original_acc_std=("original_acc", "std"),
            synthetic_acc_mean=("synthetic_acc", "mean"),
            synthetic_acc_std=("synthetic_acc", "std"),
            generalization_gap_mean=("generalization_gap", "mean"),
            train_seconds_mean=("train_seconds", "mean"),
            synthesis_seconds_mean=("synthesis_seconds", "mean"),
        )
        .reset_index()
        .sort_values(["dataset", "original_acc_mean"], ascending=[True, False])
    )
    summary_path = ROOT / "repeats_summary.csv"
    summary.to_csv(summary_path, index=False, encoding="utf-8-sig")

    sweep = load_sweeps()
    sweep_path = ROOT / "diffusion_sweep_all.csv"
    sweep.to_csv(sweep_path, index=False, encoding="utf-8-sig")

    previous = load_previous()
    previous_best = previous.loc[previous.groupby("dataset")["original_acc"].idxmax()][
        ["dataset", "method", "epsilon", "per_class", "classifier", "original_acc", "synthetic_acc"]
    ]
    new_best = sweep.loc[sweep.groupby("dataset")["original_acc"].idxmax()][
        ["dataset", "method", "epsilon", "per_class", "classifier", "original_acc", "synthetic_acc"]
    ]

    lines: list[str] = []
    lines.append("# 下一步 Diffusion+DP 实验结果汇总")
    lines.append("")
    lines.append(
        "本轮实验保持下游模型不变，仅扩展 Diffusion+DP 的 epsilon、每类合成样本数 k、训练轮数和训练超参数；GAN+DP 按计划暂不扩展。"
    )
    lines.append("")
    lines.append("## 1. 与上一轮最优结果对比")
    lines.append("")
    lines.append("| 数据集 | 上一轮最优 | 本轮扩展最优 | original_acc 提升 |")
    lines.append("|---|---:|---:|---:|")
    for dataset in ["cifar10", "mnist"]:
        old = float(previous_best[previous_best.dataset == dataset].original_acc.iloc[0])
        new = float(new_best[new_best.dataset == dataset].original_acc.iloc[0])
        lines.append(f"| {dataset.upper()} | {pct(old)} | {pct(new)} | {new * 100 - old * 100:+.2f} pp |")

    lines.append("")
    lines.append("## 2. 本轮扩展实验最优配置")
    lines.append("")
    lines.append("| 数据集 | method | epsilon | k | classifier | original_acc | synthetic_acc |")
    lines.append("|---|---|---:|---:|---|---:|---:|")
    for _, row in new_best.sort_values("dataset").iterrows():
        lines.append(
            f"| {row.dataset.upper()} | {row.method} | {row.epsilon:g} | {int(row.per_class)} | "
            f"{row.classifier} | {pct(row.original_acc)} | {pct(row.synthetic_acc)} |"
        )

    lines.append("")
    lines.append("## 3. 3 seeds 重复实验")
    lines.append("")
    lines.append("| 数据集 | config | n | original_acc mean±std | synthetic_acc mean±std | train_seconds mean |")
    lines.append("|---|---|---:|---:|---:|---:|")
    for _, row in summary.iterrows():
        lines.append(
            f"| {row.dataset.upper()} | {row.config_id} | {int(row.n)} | "
            f"{pct(row.original_acc_mean)} ± {row.original_acc_std * 100:.2f} pp | "
            f"{pct(row.synthetic_acc_mean)} ± {row.synthetic_acc_std * 100:.2f} pp | "
            f"{row.train_seconds_mean:.2f}s |"
        )

    lines.append("")
    lines.append("## 4. 结论")
    lines.append("")
    lines.append(
        "- CIFAR10 的主要收益来自 k 从 320 提升到 1000，并将 epsilon 收敛到 1.0 到 1.5 区间；单次最优达到 61.40%，3 seeds 均值约 60.22% 到 60.32%。"
    )
    lines.append(
        "- MNIST 已稳定超过 97%，最佳单次为 98.72%，重复实验均值为 98.54% ± 0.14 pp，说明该配置稳定。"
    )
    lines.append(
        "- CIFAR10 当前仍有提升空间，下一步优先围绕 `epsilon=1.0/1.25/1.5`、`k=1000/1500`、`epochs=12/15` 做小网格，同时保持 CNN/ResNet20 架构不变。"
    )
    lines.append(
        "- 可进一步尝试不改变下游模型结构的训练侧改进：CIFAR10 训练增强 RandomCrop+HorizontalFlip、label_smoothing=0.03/0.05、lr=5e-4/1e-3、weight_decay=5e-5/1e-4。"
    )

    lines.append("")
    lines.append("## 5. 输出文件")
    lines.append("")
    lines.append(f"- 明细表：`{detail_path.as_posix()}`")
    lines.append(f"- 均值标准差表：`{summary_path.as_posix()}`")
    lines.append(f"- 扩展实验总表：`{sweep_path.as_posix()}`")

    md_path = ROOT / "下一步实验结果汇总.md"
    md_path.write_text("\n".join(lines), encoding="utf-8-sig")

    report_dir = ROOT / "report"
    figure_dir = report_dir / "figures"
    figure_dir.mkdir(parents=True, exist_ok=True)

    best_compare = []
    for dataset in ["cifar10", "mnist"]:
        old = float(previous_best[previous_best.dataset == dataset].original_acc.iloc[0])
        new = float(new_best[new_best.dataset == dataset].original_acc.iloc[0])
        best_compare.append((dataset.upper(), old, new))

    fig, ax = plt.subplots(figsize=(7, 4))
    x = np.arange(len(best_compare))
    width = 0.35
    ax.bar(x - width / 2, [row[1] * 100 for row in best_compare], width, label="Previous best")
    ax.bar(x + width / 2, [row[2] * 100 for row in best_compare], width, label="New best")
    ax.set_xticks(x)
    ax.set_xticklabels([row[0] for row in best_compare])
    ax.set_ylabel("Original test accuracy (%)")
    ax.set_title("Best Result Comparison")
    ax.legend()
    ax.grid(axis="y", alpha=0.25)
    fig.tight_layout()
    best_fig = figure_dir / "best_result_comparison.png"
    fig.savefig(best_fig, dpi=200)
    plt.close(fig)

    fig, ax = plt.subplots(figsize=(8, 4))
    labels = summary["config_id"].tolist()
    means = (summary["original_acc_mean"] * 100).tolist()
    stds = (summary["original_acc_std"] * 100).tolist()
    ax.bar(np.arange(len(labels)), means, yerr=stds, capsize=5)
    ax.set_xticks(np.arange(len(labels)))
    ax.set_xticklabels(labels, rotation=20, ha="right")
    ax.set_ylabel("Original test accuracy (%)")
    ax.set_title("Three-Seed Repeat Results")
    ax.grid(axis="y", alpha=0.25)
    fig.tight_layout()
    repeat_fig = figure_dir / "repeat_mean_std.png"
    fig.savefig(repeat_fig, dpi=200)
    plt.close(fig)

    doc = Document()
    doc.add_heading("Diffusion+DP 下一步实验结果报告", level=0)
    doc.add_paragraph(
        "本报告根据最新一轮实验自动生成。实验目标是在不改变下游 CNN 与 ResNet20 模型结构的前提下，"
        "通过调节 Diffusion+DP 的 epsilon、每类合成样本数 k 与训练轮数，提高 MNIST 与 CIFAR10 的原始测试集精度。"
    )
    doc.add_heading("一、实验设置", level=1)
    doc.add_paragraph(
        "CIFAR10: epsilon=0.7/1.0/1.5/2.0, k=640/1000, epochs=10；"
        "MNIST: epsilon=0.3/0.5/0.7, k=640, epochs=10。"
        "重复实验选取 CIFAR10 前两组与 MNIST 最优组，使用 seeds=42/43/44。"
    )

    doc.add_heading("二、上一轮最优与本轮最优对比", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, text in enumerate(["数据集", "上一轮最优", "本轮扩展最优", "提升"]):
        hdr[idx].text = text
    for dataset, old, new in best_compare:
        cells = table.add_row().cells
        cells[0].text = dataset
        cells[1].text = pct(old)
        cells[2].text = pct(new)
        cells[3].text = f"{new * 100 - old * 100:+.2f} pp"
    doc.add_picture(str(best_fig), width=Inches(5.8))

    doc.add_heading("三、3 seeds 重复实验", level=1)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, text in enumerate(["数据集", "配置", "n", "original_acc mean±std", "synthetic_acc mean±std", "平均训练时间"]):
        hdr[idx].text = text
    for _, row in summary.iterrows():
        cells = table.add_row().cells
        cells[0].text = row.dataset.upper()
        cells[1].text = row.config_id
        cells[2].text = str(int(row.n))
        cells[3].text = f"{pct(row.original_acc_mean)} ± {row.original_acc_std * 100:.2f} pp"
        cells[4].text = f"{pct(row.synthetic_acc_mean)} ± {row.synthetic_acc_std * 100:.2f} pp"
        cells[5].text = f"{row.train_seconds_mean:.2f}s"
    doc.add_picture(str(repeat_fig), width=Inches(6.2))

    doc.add_heading("四、结论与下一步调参方向", level=1)
    doc.add_paragraph(
        "CIFAR10 的单次最优 original_acc 从上一轮 47.66% 提升到 61.40%，主要收益来自将 k 扩展至 1000。"
        "重复实验显示 epsilon=1.5,k=1000,CNN 的均值略高，为 60.32% ± 1.44 pp；"
        "epsilon=1.0,k=1000,CNN 的均值为 60.22% ± 1.46 pp，二者接近。"
    )
    doc.add_paragraph(
        "MNIST 已稳定超过 97%，本轮最优单次为 98.72%，重复实验均值为 98.54% ± 0.14 pp，稳定性较好。"
    )
    doc.add_paragraph(
        "下一步建议优先做 CIFAR10 小网格：epsilon=1.0/1.25/1.5，k=1000/1500，epochs=12/15；"
        "同时尝试不改变模型结构的训练侧增强，包括 RandomCrop+HorizontalFlip、label_smoothing=0.03/0.05、"
        "lr=5e-4/1e-3、weight_decay=5e-5/1e-4。"
    )

    doc_path = report_dir / "DiffusionDP_next_priority_report.docx"
    doc.save(doc_path)

    print(detail_path)
    print(summary_path)
    print(sweep_path)
    print(md_path)
    print(doc_path)


if __name__ == "__main__":
    main()
